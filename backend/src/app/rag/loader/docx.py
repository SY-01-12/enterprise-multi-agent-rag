from __future__ import annotations

import docx
from langchain_core.documents import Document

#  提取段落文本，跳过空行
def extract_paragraphs(d: docx.Document) -> list[str]:

    parts: list[str] = []
    # 遍历段落
    for para in d.paragraphs:
        # 清洗段落
        text = para.text.strip()
        # 跳过空行
        if not text:
            continue

        # 标题层级标注：Word 里自带样式：标题 1、标题 2、标题 3……，在代码里名字叫 Heading 1、Heading 2
        # 段落有样式对象，样式有 name, 样式名字是以 Heading 开头
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # 删掉 Heading → 2，再去掉空格 → 字符串"2"
            level = para.style.name.replace("Heading", "").strip()
            # 若上述取出 字符串“2” 则转为 int 且 若取不到则为 1，max() 保证最少是 1 --->  "#" * 2 ---> ## {text}
            text = f"{'#' * max(1, int(level or 1))} {text}"
        parts.append(text)
    return parts

#  提取表格内容，每行用 | 分隔
def extract_tables(d: docx.Document) -> list[str]:

    parts: list[str] = []
    # 遍历 word 中的所有表格
    for table in d.tables:
        # 遍历当前 表格中的每一行
        for row in table.rows:
            # 把一行所有单元格内容清理出来，转为列表形式
            cells = [cell.text.strip() for cell in row.cells]
            # 若不为空，把一行单元格用竖线拼接成一条字符串 如：["张三","20","男"] → "张三 | 20 | 男"
            if any(cells):
                parts.append(" | ".join(cells))
    return parts


def load_docx(file_path: str) -> list[Document]:

    # 加载文件
    d = docx.Document(file_path)

    sections: list[str] = []
    
    # 提取段落
    para_texts = extract_paragraphs(d)
    if para_texts:
        sections.append("\n".join(para_texts))
    # 提取表格
    table_texts = extract_tables(d)
    if table_texts:
        sections.append("---\n[表格内容]\n" + "\n".join(table_texts))

    # 把所有内容拼接成字符串
    full_text = "\n\n".join(sections)

    return [Document(page_content=full_text, metadata={"source": file_path})]
